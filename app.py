import streamlit as st
import pandas as pd
import copy
import re
import unicodedata
import hashlib
import time
from pathlib import Path
from datetime import datetime, date
from zoneinfo import ZoneInfo
from urllib.parse import urljoin, quote
from gspread.exceptions import APIError
from utils import (
    ensure_excel_with_sheets, append_row, update_unificado,
    PARTICIPANTES_COLS, upload_file_to_drive,
)

# Word para el documento de autorización en blanco
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Compatibilidad para versiones recientes de Streamlit donde experimental_rerun fue removido
if not hasattr(st, "experimental_rerun") and hasattr(st, "rerun"):
    st.experimental_rerun = st.rerun  # type: ignore[attr-defined]

# ===== CONFIG =====
SPREADSHEET_ID = st.secrets.get("SPREADSHEET_ID", "").strip()
BANNER_PATH = "assets/ClaveriadaBanner-1920x650.png"
SHEET_NAME = (st.secrets.get("SHEET_NAME") or "PARTICIPANTES").strip() or "PARTICIPANTES"
UPLOADS_PUBLIC_BASE_URL = (st.secrets.get("UPLOADS_PUBLIC_BASE_URL") or "").strip()
UPLOADS_DRIVE_FOLDER_ID = (
    st.secrets.get("UPLOADS_DRIVE_FOLDER_ID")
    or st.secrets.get("DRIVE_FOLDER_ID")
    or ""
).strip()
BANNER_PATH = "assets/ClaveriadaBanner-1920x650.png"
SHEET_NAME = (st.secrets.get("SHEET_NAME") or "PARTICIPANTES").strip() or "PARTICIPANTES"
UPLOADS_PUBLIC_BASE_URL = (st.secrets.get("UPLOADS_PUBLIC_BASE_URL") or "").strip()
UPLOADS_DRIVE_FOLDER_ID = (
    st.secrets.get("UPLOADS_DRIVE_FOLDER_ID")
    or st.secrets.get("DRIVE_FOLDER_ID")
    or ""
).strip()

if not SPREADSHEET_ID:
    st.error("No se encontró el ID de la hoja de cálculo en st.secrets['SPREADSHEET_ID'].")
    st.stop()

st.set_page_config(
    page_title="Claveriado RJI · Inscripción",
    layout="centered",
    initial_sidebar_state="collapsed"     # oculta la barra lateral
)

# Asegura la hoja de cálculo (no visible para usuarios) sólo una vez por sesión
_sheets_flag_key = "_sheets_initialized_for"
if st.session_state.get(_sheets_flag_key) != SPREADSHEET_ID:
    last_fail_key = "_sheets_init_last_fail"
    last_fail_ts = st.session_state.get(last_fail_key, 0.0)
    if last_fail_ts and (time.time() - float(last_fail_ts)) < 60:
        st.info("Reintentaremos conectar con la hoja de cálculo en unos segundos…")
    else:
        try:
            ensure_excel_with_sheets(SPREADSHEET_ID)
        except APIError as exc:  # type: ignore[attr-defined]
            st.session_state[last_fail_key] = time.time()
            st.warning(
                "No se pudo verificar la hoja de cálculo en este momento. "
                "Intenta nuevamente en unos minutos.\n\n"
                f"Detalle técnico: {exc}"
            )
        else:
            st.session_state[_sheets_flag_key] = SPREADSHEET_ID
            st.session_state.pop(last_fail_key, None)

# ===== Estilos (paleta Claveriada + ocultar sidebar) =====
st.markdown(
    """
    <style>
    /* Ocultar sidebar y el botón de despliegue */
    [data-testid="stSidebar"], [data-testid="collapsedControl"] { display: none !important; }

    :root{
      --bg:#141d2c;           /* fondo principal */
      --card:#1b2a44;         /* tarjetas */
      --border:#243656;       /* bordes */
      --text:#e7eefc;         /* texto principal */
      --muted:#9fb1d0;        /* texto secundario */
      --accent:#ff9c2a;       /* naranja */
      --accent2:#8bd143;      /* verde */
      --accent3:#9cc5ff;      /* azul claro */
      --radius:18px;
    }
    .main{ background:var(--bg); }
    .block-container{ padding-top:1rem; padding-bottom:3rem; max-width:980px; }
    .rji-card{
      background:var(--card); border:1px solid var(--border);
      border-radius:var(--radius); padding:1.25rem 1.5rem;
      box-shadow:0 16px 40px rgba(0,0,0,.35);
    }
    .rji-title{ color:var(--text); font-size:2rem; font-weight:800; margin:.2rem 0 .2rem; }
    .rji-sub{ color:var(--muted); margin-bottom:1.2rem; font-size:.98rem; }
    .stTabs [data-baseweb="tab-list"]{ gap:6px; }
    .stTabs [data-baseweb="tab"]{
      background:var(--card); border-radius:999px; padding:.45rem .95rem;
      border:1px solid var(--border); color:var(--text);
    }
    .stTabs [aria-selected="true"]{
      border:1px solid var(--accent); background:#233658; color:var(--accent);
    }
    .stage-progress{ margin:1.75rem 0 0; padding-top:1rem; border-top:1px solid rgba(255,255,255,.12); }
    .stage-progress-label{ display:flex; justify-content:space-between; font-weight:600; color:var(--text); margin-bottom:.35rem; }
    .stage-progress-bar{ background:rgba(255,255,255,.2); border-radius:999px; height:20px; overflow:hidden; border:1px solid var(--border); box-shadow:0 0 12px rgba(255,156,42,.35) inset; }
    .stage-progress-bar span{ display:block; height:100%; background:linear-gradient(90deg,var(--accent),var(--accent3)); border-radius:inherit; transition:width .35s ease; }
    .stage-progress-sub{ color:var(--text); font-size:.95rem; margin-top:.45rem; font-weight:500; }
    .motivacion-box{ background:rgba(255,156,42,.12); border:1px solid rgba(255,156,42,.35); color:var(--accent); padding:.75rem 1rem; border-radius:12px; font-weight:600; }
    .perfil-slider-labels{ display:flex; justify-content:space-between; color:var(--muted); font-weight:600; margin-top:.35rem; }
    /* Inputs */
    label, .stMarkdown, .stCaption, .stRadio, .stText, .stSelectbox, .stDateInput, .stTimeInput{
      color:var(--text) !important;
    }
    .stTextInput>div>div>input, .stTextArea textarea{
      background:var(--bg) !important; color:var(--text) !important; border-radius:10px;
      border:1px solid var(--border) !important;
    }
    .stSelectbox>div>div{ background:var(--bg) !important; border-radius:10px; border:1px solid var(--border); }
    .stDateInput>div>div, .stTimeInput>div>div{
      background:var(--bg) !important; border:1px solid var(--border) !important; border-radius:10px;
    }
    /* Botones */
    .stButton>button, .stDownloadButton>button{
      background:var(--accent); color:#1a1625; border:1px solid var(--accent);
      border-radius:10px; padding:.5rem 1rem; font-weight:700;
    }
    .stButton>button:hover, .stDownloadButton>button:hover{ filter:brightness(1.05); }

    /* Banner con bordes redondeados */
    .banner-wrap img{ border-radius:16px; border:1px solid var(--border); }
    </style>
    """,
    unsafe_allow_html=True
)

# ===== Banner =====
try:
    st.markdown('<div class="banner-wrap">', unsafe_allow_html=True)
    st.image(BANNER_PATH, use_column_width=True)
    st.markdown("</div>", unsafe_allow_html=True)
except Exception:
    pass
st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)

# ===== Header =====
st.markdown('<div class="rji-card">', unsafe_allow_html=True)
st.markdown('<div class="rji-title">Inscripciones · RJI</div>', unsafe_allow_html=True)
st.markdown('<div class="rji-sub">Participantes y Acompañantes — Medellín, Colombia</div>', unsafe_allow_html=True)

# ===== Pestañas =====
tab1, tab2, tab3 = st.tabs(["Participante", "Acompañante/Institución", "Voluntarios"])

# ===== Utilidades =====
def crear_doc_autorizacion_en_blanco(logo_path="assets/logo.png"):
    """Documento Word en blanco (no depende de datos)."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2); section.bottom_margin = Cm(2); section.left_margin = Cm(2); section.right_margin = Cm(2)
    try:
        if Path(logo_path).exists():
            header = doc.sections[0].header
            hdr_p = header.paragraphs[0]
            run = hdr_p.add_run()
            run.add_picture(logo_path, width=Inches(3.0))
    except Exception:
        pass
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FORMATO DE AUTORIZACIÓN Y ACOMPAÑAMIENTO"); r.bold = True; r.font.size = Pt(16)
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run("Claveriada RJI – Medellín, Colombia").font.size = Pt(12)
    info = doc.add_paragraph()
    info.add_run("La información consignada es sensible y será utilizada únicamente para construir el perfil de cada participante "
                 "en la Claveriada RJI y para la logística del encuentro. No será compartida con terceros.\n").font.size = Pt(10)
    a = doc.add_paragraph()
    a.add_run("Yo, ________________________________, identificado(a) con documento No. ________________________, "
              "en calidad de acudiente/acompañante, autorizo la participación de las/los siguientes jóvenes en el evento.").font.size = Pt(11)
    t = doc.add_table(rows=2, cols=2); t.style = "Table Grid"
    t.cell(0,0).text = "Correo del acompañante"; t.cell(0,1).text = "_______________________________"
    t.cell(1,0).text = "Teléfono del acompañante"; t.cell(1,1).text = "_______________________________"
    doc.add_paragraph().add_run("Relación de jóvenes a cargo").bold = True
    tbl = doc.add_table(rows=1, cols=5); tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Nombre completo"; hdr[1].text = "Documento"; hdr[2].text = "Edad"; hdr[3].text = "EPS"; hdr[4].text = "Complicaciones de salud"
    for _ in range(6):
        row = tbl.add_row().cells
        for i in range(5): row[i].text = ""
    doc.add_paragraph("Declaro que la información es veraz y me comprometo a acompañar y velar por el bienestar de las/los jóvenes, "
                      "cumplir las indicaciones del equipo organizador y notificar cualquier situación de salud o emergencia.")
    f = doc.add_table(rows=2, cols=2); f.autofit = True
    f.cell(0,0).text = "\n\n_______________________________"
    f.cell(0,1).text = "\n\n_______________________________"
    f.cell(1,0).text = "Firma del acompañante"
    f.cell(1,1).text = "Firma de la institución"
    return doc

def calcular_edad(fecha_str):
    if not fecha_str:
        return ""
    try:
        d = pd.to_datetime(str(fecha_str), errors="coerce")
        if pd.isna(d): return ""
        today = pd.Timestamp.today().date()
        d = d.date()
        return today.year - d.year - ((today.month, today.day) < (d.month, d.day))
    except Exception:
        return ""


def _clean_string(value: object) -> str:
    """Return a normalized string suitable for storage without mutating widgets."""
    if not isinstance(value, str):
        return ""
    normalized = unicodedata.normalize("NFKC", value)
    normalized = normalized.strip()
    # Collapse any internal whitespace runs to a single space
    normalized = re.sub(r"\s+", " ", normalized)
    return normalized


def _clean_phone_number(value: object) -> str:
    """Normalize a phone number keeping digits and an optional leading plus."""
    if not isinstance(value, str):
        return ""
    normalized = unicodedata.normalize("NFKC", value)
    normalized = normalized.strip()
    if not normalized:
        return ""
    has_plus = normalized.startswith("+")
    digits = "".join(ch for ch in normalized if ch.isdigit())
    if not digits:
        return ""
    return f"+{digits}" if has_plus else digits


def _format_phone_for_sheet(value: str) -> str:
    """Prefix phone numbers with ' so Sheets treats them as literal text."""
    if not value:
        return ""
    return f"'{value}"


def _format_upload_for_sheet(path_value: str, display_name: str = "") -> str:
    """Return a Sheets-friendly value (hyperlink when possible) for uploaded files."""
    if not path_value:
        return ""

    path_value = str(path_value)
    if path_value.startswith("="):
        return path_value

    filename = display_name or Path(path_value).name or "Archivo"
    safe_label = filename.replace('"', '""')

    if path_value.startswith(("http://", "https://")):
        url = path_value.replace('"', "%22")
        return f'=HYPERLINK("{url}", "{safe_label}")'

    public_base = UPLOADS_PUBLIC_BASE_URL.rstrip("/")
    if public_base:
        relative = Path(path_value).as_posix()
        # Remove leading relative markers and root slashes to avoid urljoin resets.
        while relative.startswith("./"):
            relative = relative[2:]
        if relative.startswith("/"):
            relative = relative[1:]
        if relative.startswith("uploads/"):
            relative_fragment = relative[len("uploads/"):]
        else:
            relative_fragment = relative
        base = f"{public_base}/"
        relative_encoded = quote(relative_fragment, safe="/")
        url = urljoin(base, relative_encoded)
        url = url.replace('"', "%22")
        return f'=HYPERLINK("{url}", "{safe_label}")'

    return path_value


def _get_participant_payload() -> dict:
    """Centralized storage for the participant export row."""
    return st.session_state.setdefault("_participant_payload", {})

COLOMBIA_DEPARTAMENTOS = [
    "Amazonas", "Antioquia", "Arauca", "Atlántico", "Bogotá D.C.", "Bolívar", "Boyacá", "Caldas",
    "Caquetá", "Casanare", "Cauca", "Cesar", "Chocó", "Córdoba", "Cundinamarca", "Guainía",
    "Guaviare", "Huila", "La Guajira", "Magdalena", "Meta", "Nariño", "Norte de Santander",
    "Putumayo", "Quindío", "Risaralda", "San Andrés y Providencia", "Santander", "Sucre", "Tolima",
    "Valle del Cauca", "Vaupés", "Vichada"
]

COLOMBIA_CIUDADES = {
    "Amazonas": ["Leticia", "Puerto Nariño"],
    "Antioquia": ["Medellín", "Bello", "Envigado", "Itagüí", "Rionegro", "Apartadó", "Turbo", "La Ceja", "Caucasia", "Sabaneta"],
    "Arauca": ["Arauca", "Saravena", "Tame"],
    "Atlántico": ["Barranquilla", "Soledad", "Malambo", "Sabanalarga", "Puerto Colombia"],
    "Bogotá D.C.": ["Bogotá"],
    "Bolívar": ["Cartagena", "Magangué", "Turbaco", "Arjona", "El Carmen de Bolívar"],
    "Boyacá": ["Tunja", "Duitama", "Sogamoso", "Chiquinquirá", "Paipa"],
    "Caldas": ["Manizales", "Villamaría", "Chinchiná", "La Dorada"],
    "Caquetá": ["Florencia", "San Vicente del Caguán", "Belén de los Andaquíes"],
    "Casanare": ["Yopal", "Aguazul", "Villanueva", "Tauramena"],
    "Cauca": ["Popayán", "Santander de Quilichao", "Puerto Tejada", "Guapi"],
    "Cesar": ["Valledupar", "Aguachica", "La Jagua de Ibirico", "Bosconia"],
    "Chocó": ["Quibdó", "Istmina", "Tadó", "Condoto"],
    "Córdoba": ["Montería", "Lorica", "Sahagún", "Planeta Rica", "Tierralta"],
    "Cundinamarca": ["Soacha", "Chía", "Zipaquirá", "Facatativá", "Girardot"],
    "Guainía": ["Inírida"],
    "Guaviare": ["San José del Guaviare", "Calamar"],
    "Huila": ["Neiva", "Pitalito", "Garzón", "La Plata"],
    "La Guajira": ["Riohacha", "Maicao", "Uribia", "Fonseca"],
    "Magdalena": ["Santa Marta", "Ciénaga", "Fundación", "El Banco"],
    "Meta": ["Villavicencio", "Acacías", "Granada", "Puerto López"],
    "Nariño": ["Pasto", "Ipiales", "Tumaco", "Túquerres"],
    "Norte de Santander": ["Cúcuta", "Ocaña", "Pamplona", "Villa del Rosario"],
    "Putumayo": ["Mocoa", "Puerto Asís", "Villagarzón"],
    "Quindío": ["Armenia", "Circasia", "Montenegro", "Quimbaya"],
    "Risaralda": ["Pereira", "Dosquebradas", "Santa Rosa de Cabal", "La Virginia"],
    "San Andrés y Providencia": ["San Andrés", "Providencia"],
    "Santander": ["Bucaramanga", "Floridablanca", "Girón", "Barrancabermeja", "San Gil"],
    "Sucre": ["Sincelejo", "Corozal", "Tolú", "San Marcos"],
    "Tolima": ["Ibagué", "Espinal", "Honda", "Melgar"],
    "Valle del Cauca": ["Cali", "Palmira", "Buenaventura", "Buga", "Tuluá", "Yumbo"],
    "Vaupés": ["Mitú"],
    "Vichada": ["Puerto Carreño", "La Primavera"],
}

def _todas_las_ciudades():
    ciudades = []
    for lista in COLOMBIA_CIUDADES.values():
        ciudades.extend(lista)
    return sorted(set(ciudades))

TODAS_LAS_CIUDADES = _todas_las_ciudades()

OBRAS_RJI = [
    "Bethlemitas",
    "Casa de Ejercicios San Ignacio de Pasto",
    "Casa Pastoral Nuestra Señora del Camino",
    "CIJ - Casa Ignaciana de la Juventud",
    "Colegio Berchmans",
    "Colegio Fe y Alegría Antonio Nariño",
    "Colegio Fe y Alegría Colorados",
    "Colegio Fe y Alegría Gabriel García Marquez",
    "Colegio Fe y Alegría Germán Vargas Cantillo",
    "Colegio Fe y Alegría IED Las Mercedes",
    "Colegio Fe y Alegría IED Torquigua",
    "Colegio Fe y Alegría Jaime Salazar",
    "Colegio Fe y Alegría Jose Maria Velaz",
    "Colegio Fe y Alegría Jose Raimundo Sojo",
    "Colegio Fe y Alegría Juan Francisco Sarasti",
    "Colegio Fe y Alegría Libertador Simón Bolivar",
    "Colegio Fe y Alegría Luis Felipe Cabrera",
    "Colegio Fe y Alegría Mario Moreno",
    "Colegio Fe y Alegría Mesetas de San Rafael",
    "Colegio Fe y Alegría Nuestra Señora de Belén",
    "Colegio Fe y Alegría San Ignacio",
    "Colegio Fe y Alegría San Vicente",
    "Colegio Reina de la Paz",
    "Colegio San Bartolomé la Merced",
    "Colegio San Francisco Javier",
    "Colegio San Ignacio de Medellín",
    "Colegio San José",
    "Colegio San Luis Gonzaga",
    "Colegio San Pedro Claver",
    "Colegio Santa Luisa",
    "Escuela de Fútbol Fundación Juan Lorenzo Lucero",
    "Fundación Suyusama",
    "IMCA - Instituto Mayor Campesino",
    "Institucion educativa Antonio Nariño",
    "Misión San Pedro Claver Cartagena",
    "Museo Juan Lorenzo Lucero",
    "Pontificia Universidad Javeriana Bogotá",
    "Pontificia Universidad Javeriana Cali",
    "RJI - Red Juvenil Ignaciana",
    "Templo de Cristo Rey",
]

PARENTESCOS = [
    "Madre", "Padre", "Hermana / Hermano", "Tía / Tío", "Abuela / Abuelo", "Tutor/a legal", "Acompañante de obra", "Otro"
]


PARTICIPANT_DEFAULTS = {
    "part_step": 1,
    "part_es_mayor_option": "",
    "part_tipo_doc_p": "",
    "part_doc_p": "",
    "_clean_part_doc_p": "",
    "part_doc_id_name": "",
    "part_doc_id_bytes": b"",
    "part_nombres": "",
    "part_apellidos": "",
    "part_apodo": "",
    "part_tel": "",
    "part_correo": "",
    "part_direccion": "",
    "part_region": "",
    "part_ciudad": "",
    "part_fecha_nac": date(2006, 1, 1),
    "part_talla": "",
    "part_eps": "",
    "part_rest_alim": "",
    "part_salud_mental": "",
    "part_obra": "",
    "part_obra_select": "",
    "part_obra_custom": "",
    "part_proceso": "",
    "part_tipo_doc_a": "",
    "part_doc_a": "",
    "_clean_part_doc_a": "",
    "part_nom_a": "",
    "part_ape_a": "",
    "part_correo_a": "",
    "part_tel_a": "",
    "part_parentesco_a": "",
    "part_exp_sig": "",
    "part_intereses": [],
    "part_dato_freak": "",
    "part_pregunta": "",
    "part_motivo": "",
    "part_preguntas_frec": "",
    "part_acomp_amigos": False,
    "part_acomp_familia": False,
    "part_acomp_escucha": False,
    "part_acomp_mentoria": False,
    "part_acomp_espiritual": False,
    "part_acomp_red_comunidad": False,
    "part_acomp_ninguna": False,
    "part_conoce_rji": "",
    "part_perfil_slider": 1,
    "part_acepta_datos": False,
    "part_acepta_whatsapp": False,
}


def _init_participant_state():
    for key, value in PARTICIPANT_DEFAULTS.items():
        if key not in st.session_state:
            st.session_state[key] = copy.deepcopy(value)


def _reset_participant_state():
    """Reset stored participant answers after a successful submission."""
    for key in PARTICIPANT_DEFAULTS:
        st.session_state.pop(key, None)

    # Clear transient helper caches so toggles and uploads start fresh.
    for transient_key in (
        "_prev_part_acomp_values",
        "_prev_part_acomp_ninguna",
        "part_exp_order",
        "exp_sort",
        "part_doc_archivo",
        "_part_doc_drive_hash",
        "_part_doc_drive_link",
    ):
        st.session_state.pop(transient_key, None)

    # Remove legacy keys from sesiones previas.
    for legacy_key in (
        "part_tipo_doc_a",
        "part_doc_a",
        "_clean_part_doc_a",
        "part_contact_doc",
        "part_contact_doc_name",
        "part_contact_doc_bytes",
        "_contact_doc_drive_hash",
        "_contact_doc_drive_link",
    ):
        st.session_state.pop(legacy_key, None)

    # Remove legacy keys from sesiones previas.
    for legacy_key in (
        "part_contact_doc_name",
        "part_contact_doc_bytes",
    ):
        st.session_state.pop(legacy_key, None)

    # Remove legacy keys from sesiones previas.
    for legacy_key in (
        "part_contact_doc_name",
        "part_contact_doc_bytes",
    ):
        st.session_state.pop(legacy_key, None)

    st.session_state.pop("_participant_payload", None)
    st.session_state.pop("_participant_reset_pending", None)


def _participant_stage_fields(stage: int):
    base = {
        1: [
            "part_es_mayor_option", "part_tipo_doc_p", "part_doc_p", "part_nombres",
            "part_apellidos", "part_apodo", "part_tel", "part_correo", "part_direccion",
            "part_region", "part_ciudad", "part_fecha_nac", "part_talla", "part_eps",
            "part_rest_alim", "part_salud_mental", "part_obra", "part_proceso",
            "part_nom_a", "part_ape_a", "part_parentesco_a", "part_tel_a",
        ],
        2: [
            "part_exp_sig", "part_intereses", "part_dato_freak", "part_pregunta",
        ],
        3: [
            "part_motivo", "part_preguntas_frec",
            "part_acomp_amigos", "part_acomp_familia", "part_acomp_escucha", "part_acomp_mentoria",
            "part_acomp_espiritual", "part_acomp_red_comunidad", "part_acomp_ninguna", "part_conoce_rji",
            "part_acepta_datos",
        ],
    }
    if stage == 1:
        base[1].append("part_correo_a")
    return base.get(stage, [])


def _value_is_filled(val, key: str) -> bool:
    optional_blanks = {"part_apodo", "part_salud_mental", "part_preguntas_frec", "part_correo_a"}
    if key in optional_blanks:
        return True
    if key == "part_acepta_datos":
        return bool(val)
    if isinstance(val, str):
        return bool(val.strip())
    if isinstance(val, list):
        return len(val) > 0
    if isinstance(val, bool):
        if key.startswith("part_acomp_"):
            return True
        return bool(val)
    if isinstance(val, (date, datetime)):
        if key == "part_fecha_nac" and isinstance(val, date) and val == PARTICIPANT_DEFAULTS["part_fecha_nac"]:
            return False
        return True
    if val is None:
        return False
    return True


def _stage_progress(stage: int):
    fields = _participant_stage_fields(stage)
    if not fields:
        return 0.0, 0, 0, 0
    answered = sum(1 for key in fields if _value_is_filled(st.session_state.get(key), key))
    total = len(fields)
    return answered / total, int(round((answered / total) * 100)), answered, total


def _goto_participant_stage(stage: int):
    st.session_state.part_step = stage


def _emit_stage_errors(messages, show: bool = True) -> bool:
    if show:
        for message in messages:
            st.error(message)
    return len(messages) == 0


def _normalize_numeric_input(value: str) -> tuple[bool, str]:
    raw = (value or "").strip()
    if not raw:
        return False, ""
    allowed = set("0123456789 .-")
    if any(ch not in allowed for ch in raw):
        return False, ""
    digits = "".join(ch for ch in raw if ch.isdigit())
    if not digits:
        return False, ""
    return True, digits


def _validate_participant_stage1(show_errors: bool = True) -> bool:
    errors = []
    mayor_option = st.session_state.get("part_es_mayor_option", "")
    if mayor_option not in {"Sí", "No"}:
        errors.append("Confírmanos si eres mayor de edad para continuar.")

    doc_ok, cleaned_doc = _normalize_numeric_input(st.session_state.get("part_doc_p", ""))
    if doc_ok:
        st.session_state["_clean_part_doc_p"] = cleaned_doc
    else:
        st.session_state["_clean_part_doc_p"] = ""
        errors.append("El documento del participante debe contener solo dígitos.")

    if not st.session_state.get("part_tipo_doc_p"):
        errors.append("Selecciona el tipo de documento del participante.")

    if not st.session_state.get("part_nombres", "").strip() or not st.session_state.get("part_apellidos", "").strip():
        errors.append("Ingresa tus nombres y apellidos tal como aparecen en tu documento.")

    if not st.session_state.get("part_direccion", "").strip():
        errors.append("Cuéntanos tu dirección de residencia.")

    if not st.session_state.get("part_region") or not st.session_state.get("part_ciudad"):
        errors.append("Selecciona tu región y ciudad para continuar.")

    if not st.session_state.get("part_talla"):
        errors.append("Selecciona tu talla de camiseta.")

    tel_clean = _clean_phone_number(st.session_state.get("part_tel", ""))
    if not tel_clean:
        errors.append("Déjanos un número de contacto personal válido (puedes incluir el prefijo +57).")

    if not st.session_state.get("part_correo", "").strip():
        errors.append("Incluye un correo de contacto personal.")

    nom_a = st.session_state.get("part_nom_a", "").strip()
    ape_a = st.session_state.get("part_ape_a", "").strip()
    tel_a_clean = _clean_phone_number(st.session_state.get("part_tel_a", ""))
    parentesco = st.session_state.get("part_parentesco_a", "")

    contacto_name_issue = not nom_a or not ape_a
    es_menor = mayor_option == "No"
    if es_menor and contacto_name_issue:
        errors.append("Para menores, los nombres y apellidos del acudiente son obligatorios.")
    elif contacto_name_issue:
        errors.append("Ingresa nombres y apellidos del contacto de emergencia.")

    if not tel_a_clean:
        errors.append("Incluye un teléfono válido para el contacto de emergencia (puedes incluir el prefijo +57).")

    if not parentesco:
        errors.append("Selecciona el parentesco o vínculo del contacto de emergencia.")

    es_mayor_bool = None
    if mayor_option == "Sí":
        es_mayor_bool = True
    elif mayor_option == "No":
        es_mayor_bool = False

    is_valid = len(errors) == 0
    if is_valid:
        payload = _get_participant_payload()
        payload.update(
            {
                "es_mayor_edad": es_mayor_bool,
                "tipo_documento_participante": st.session_state.get("part_tipo_doc_p", ""),
                "documento_participante": cleaned_doc if doc_ok else _clean_string(st.session_state.get("part_doc_p", "")),
                "nombres": _clean_string(st.session_state.get("part_nombres", "")),
                "apellidos": _clean_string(st.session_state.get("part_apellidos", "")),
                "como_te_gusta_que_te_digan": _clean_string(st.session_state.get("part_apodo", "")),
                "telefono_celular": tel_clean,
                "correo": _clean_string(st.session_state.get("part_correo", "")),
                "direccion": _clean_string(st.session_state.get("part_direccion", "")),
                "region": _clean_string(st.session_state.get("part_region", "")),
                "ciudad": _clean_string(st.session_state.get("part_ciudad", "")),
                "fecha_nacimiento": st.session_state.get("part_fecha_nac"),
                "talla_camisa": st.session_state.get("part_talla", ""),
                "eps": _clean_string(st.session_state.get("part_eps", "")),
                "restricciones_alimentarias": _clean_string(st.session_state.get("part_rest_alim", "")),
                "salud_mental": _clean_string(st.session_state.get("part_salud_mental", "")),
                "obra_institucion": _clean_string(st.session_state.get("part_obra", "")),
                "proceso_juvenil": _clean_string(st.session_state.get("part_proceso", "")),
                "tipo_documento_contacto": "",
                "documento_contacto": "",
                "nombres_contacto": _clean_string(nom_a),
                "apellidos_contacto": _clean_string(ape_a),
                "telefono_contacto": tel_a_clean,
                "correo_contacto": _clean_string(st.session_state.get("part_correo_a", "")),
                "parentesco_contacto": _clean_string(parentesco),
            }
        )

    return _emit_stage_errors(errors, show_errors)


def _validate_participant_stage2(show_errors: bool = True) -> bool:
    errors = []
    if not st.session_state.get("part_exp_sig", "").strip():
        errors.append("Cuéntanos una experiencia juvenil significativa.")

    intereses = st.session_state.get("part_intereses", [])
    if not intereses:
        errors.append("Selecciona al menos un interés personal (hasta 3).")

    if not st.session_state.get("part_dato_freak", "").strip():
        errors.append("Comparte un hobby o dato curioso para continuar.")

    if not st.session_state.get("part_pregunta", "").strip():
        errors.append("Propón una pregunta para conectar con otros participantes.")
    is_valid = len(errors) == 0
    if is_valid:
        payload = _get_participant_payload()
        payload.update(
            {
                "experiencia_significativa": _clean_string(st.session_state.get("part_exp_sig", "")),
                "intereses_personales": list(st.session_state.get("part_intereses", [])),
                "hobby_o_dato_curioso": _clean_string(st.session_state.get("part_dato_freak", "")),
                "pregunta_para_conectar": _clean_string(st.session_state.get("part_pregunta", "")),
            }
        )

    return _emit_stage_errors(errors, show_errors)


def _validate_participant_stage3(show_errors: bool = True) -> bool:
    errors = []
    if not st.session_state.get("part_motivo", "").strip():
        errors.append("Cuéntanos por qué te interesa tu experiencia prioritaria.")

    if st.session_state.get("part_conoce_rji") == "":
        errors.append("Cuéntanos si conoces la RJI antes de guardar.")

    if not st.session_state.get("part_acepta_datos"):
        errors.append("Debes aceptar el aviso de privacidad.")

    if not st.session_state.get("part_acepta_whatsapp"):
        errors.append("Debes autorizar la comunicación por WhatsApp.")

    return _emit_stage_errors(errors, show_errors)


def render_stage_progress(stage: int):
    _, porcentaje, respondidas, total = _stage_progress(stage)
    st.markdown(
        f"""
        <div class=\"stage-progress\">
            <div class=\"stage-progress-label\">
                <span>Avance de la etapa</span>
                <span>{porcentaje}%</span>
            </div>
            <div class=\"stage-progress-bar\">
                <span style=\"width:{porcentaje}%\"></span>
            </div>
        </div>
        <div class=\"stage-progress-sub\">Has respondido {respondidas} de {total} preguntas clave en esta sección.</div>
        """,
        unsafe_allow_html=True,
    )


if st.session_state.get("_participant_reset_pending"):
    _reset_participant_state()

_init_participant_state()

# ================= PARTICIPANTE =================
with tab1:
    success_message = st.session_state.pop("_participant_success_message", "")
    if success_message:
        st.success(success_message)

    stage = st.session_state.part_step
    stage_titles = {
        1: "Etapa 1 · Datos personales",
        2: "Etapa 2 · Historial e intereses",
        3: "Etapa 3 · Experiencias y acompañamiento",
    }
    motivaciones = {
        1: "Vamos paso a paso, comparte quién eres para arrancar con buen pie 💪",
        2: "¡Bien! Ya casi llegamos a las experiencias, cuéntanos lo que te mueve ✨",
        3: "Último tramo, vamos con toda para elegir experiencias y acompañamientos 🚀",
    }

    st.markdown(f"### {stage_titles.get(stage, '')}")
    st.markdown(f"<div class='motivacion-box'>{motivaciones.get(stage, '')}</div>", unsafe_allow_html=True)

    experiencias = ["Servicio", "Peregrinaje", "Cultura y arte", "Espiritualidad", "Vocación", "Incidencia política"]

    ACOMP_KEYS = [
        "part_acomp_familia",
        "part_acomp_amigos",
        "part_acomp_escucha",
        "part_acomp_mentoria",
        "part_acomp_espiritual",
        "part_acomp_red_comunidad",
    ]

    if stage == 1:
        with st.form("form_participante_stage1", clear_on_submit=False):
            st.subheader("Información básica")
            mayor_options = ["", "Sí", "No"]
            current_mayor = st.session_state.get("part_es_mayor_option", "")
            if current_mayor not in mayor_options:
                st.session_state.part_es_mayor_option = ""
            st.selectbox(
                "¿Eres mayor de edad?",
                mayor_options,
                key="part_es_mayor_option",
                format_func=lambda val: "Selecciona una opción" if val == "" else val,
            )

            doc_options = ["", "CC", "TI", "CE", "Pasaporte", "Otro"]
            current_doc = st.session_state.get("part_tipo_doc_p", "")
            if current_doc not in doc_options:
                st.session_state.part_tipo_doc_p = ""
            st.selectbox(
                "Tipo de documento",
                doc_options,
                key="part_tipo_doc_p",
                format_func=lambda val: "Selecciona el tipo de documento" if val == "" else val,
            )
            st.text_input("Número de documento (solo dígitos)", max_chars=20, placeholder="Ej: 1234567890", key="part_doc_p")
            doc_file = st.file_uploader(
                "Adjunta copia del documento (PDF o imagen)",
                type=["pdf", "png", "jpg", "jpeg"],
                key="part_doc_archivo",
            )
            if doc_file is not None:
                st.session_state.part_doc_id_name = doc_file.name
                st.session_state.part_doc_id_bytes = doc_file.getbuffer().tobytes()
            elif st.session_state.get("part_doc_id_name"):
                st.caption(f"Archivo guardado: {st.session_state.get('part_doc_id_name')}")

            nombres = st.text_input("Nombres", placeholder="Como aparecen en tu documento", key="part_nombres")
            apellidos = st.text_input("Apellidos", placeholder="Como aparecen en tu documento", key="part_apellidos")
            apodo = st.text_input("¿Cómo te gusta que te digan?", placeholder="Opcional", key="part_apodo")
            telefono = st.text_input("Teléfono celular", placeholder="+57 ...", key="part_tel")
            correo = st.text_input("Correo", placeholder="tu@correo.com", key="part_correo")
            direccion = st.text_input("Dirección de residencia", placeholder="Barrio, calle, número", key="part_direccion")

            col_reg, col_ciudad = st.columns(2)
            region_options = [""] + COLOMBIA_DEPARTAMENTOS
            current_region = st.session_state.get("part_region", "")
            if current_region not in region_options:
                st.session_state.part_region = ""
            col_reg.selectbox(
                "Región / Departamento",
                region_options,
                key="part_region",
                format_func=lambda val: "Selecciona la región / departamento" if val == "" else val,
            )

            ciudad_options = [""] + TODAS_LAS_CIUDADES
            current_ciudad = st.session_state.get("part_ciudad", "")
            if current_ciudad not in ciudad_options:
                st.session_state.part_ciudad = ""
            col_ciudad.selectbox(
                "Ciudad / Municipio",
                ciudad_options,
                key="part_ciudad",
                format_func=lambda val: "Selecciona la ciudad / municipio" if val == "" else val,
            )

            st.date_input(
                "Fecha de nacimiento",
                min_value=date(1900, 1, 1),
                max_value=date.today(),
                key="part_fecha_nac"
            )
            talla_options = ["", "16", "XS", "S", "M", "L", "XL", "2XL"]
            current_talla = st.session_state.get("part_talla", "")
            if current_talla not in talla_options:
                st.session_state.part_talla = ""
            st.selectbox(
                "Talla de camiseta",
                talla_options,
                key="part_talla",
                format_func=lambda val: "Selecciona tu talla" if val == "" else val,
            )

            eps = st.text_input("EPS", placeholder="Escribe tu EPS", key="part_eps")
            rest_alim = st.text_input(
                "Restricciones alimentarias (o 'ninguna')",
                placeholder="Vegetariano, alergias, etc.",
                key="part_rest_alim"
            )
            salud = st.text_area(
                "Complicaciones/alertas de salud (solo lo necesario para cuidarte mejor)",
                key="part_salud_mental"
            )

            obra_options = [""] + OBRAS_RJI + ["Otra / No aparece en la lista"]
            current_select = st.session_state.get("part_obra_select", "")
            if current_select not in obra_options:
                st.session_state.part_obra_select = ""
            st.selectbox(
                "¿De qué obra / institución vienes?",
                obra_options,
                key="part_obra_select",
                format_func=lambda val: "Selecciona la obra o institución" if val == "" else val,
            )
            obra_sel = st.session_state.get("part_obra_select", "")
            if obra_sel == "Otra / No aparece en la lista":
                otra_obra = st.text_input(
                    "Escribe el nombre de tu obra / institución",
                    key="part_obra_custom"
                )
                st.session_state.part_obra = otra_obra.strip()
            elif obra_sel == "":
                st.session_state.part_obra = ""
            else:
                st.session_state.part_obra = obra_sel

            proceso = st.text_input(
                "¿Perteneces a algún proceso juvenil? ¿Cuál?",
                placeholder="Nombre del proceso",
                key="part_proceso"
            )

            st.subheader("Contacto de emergencia / acudiente")
            st.caption("Incluye la persona que estará disponible ante cualquier emergencia.")
            doc_ac_options = ["", "CC", "CE", "Pasaporte", "Otro"]
            current_doc_a = st.session_state.get("part_tipo_doc_a", "")
            if current_doc_a not in doc_ac_options:
                st.session_state.part_tipo_doc_a = ""
            st.text_input("Nombres del contacto", key="part_nom_a")
            st.text_input("Apellidos del contacto", key="part_ape_a")
            st.selectbox(
                "Tipo de documento (contacto)",
                doc_ac_options,
                key="part_tipo_doc_a",
                format_func=lambda val: "Selecciona el documento" if val == "" else val,
            )
            st.text_input(
                "Documento del contacto (solo dígitos)",
                max_chars=20,
                placeholder="Ej: 1012345678",
                key="part_doc_a"
            )
            st.text_input("Teléfono del contacto", key="part_tel_a")
            st.text_input("Correo del contacto (opcional)", key="part_correo_a")
            parentesco_opciones = [""] + PARENTESCOS
            current_parentesco = st.session_state.get("part_parentesco_a", "")
            if current_parentesco not in parentesco_opciones:
                st.session_state.part_parentesco_a = ""
            st.selectbox(
                "Parentesco o vínculo",
                parentesco_opciones,
                key="part_parentesco_a",
                format_func=lambda val: "Selecciona el parentesco" if val == "" else val,
            )
            avanzar = st.form_submit_button("Avanzar a intereses", use_container_width=True)
            if avanzar:
                if _validate_participant_stage1():
                    _goto_participant_stage(2)

    elif stage == 2:
        intereses_full = [
            "Aventura", "Deporte", "Contemplación", "Arte", "Música", "Danza", "Teatro", "Fotografía",
            "Ciencia", "Tecnología", "Videojuegos", "Cocina", "Emprendimiento", "Lectura", "Naturaleza",
            "Ecología integral", "Montaña", "Ciclismo", "Senderismo", "Viajes", "Idiomas",
            "Servicio comunitario", "Liderazgo", "Mascotas"
        ]
        with st.form("form_participante_stage2", clear_on_submit=False):
            st.subheader("Momentos que te han marcado")
            st.text_area(
                "Experiencia juvenil significativa (torneo, voluntariado, congreso, etc.)",
                key="part_exp_sig"
            )
            st.multiselect(
                "Intereses personales",
                intereses_full,
                key="part_intereses",
                max_selections=3,
                help="Selecciona hasta 3 intereses que hoy te representen más."
            )
            st.text_input("Hobby o dato curioso que quieras compartir", placeholder="Algo que te represente", key="part_dato_freak")
            st.text_input("Propón una pregunta para conectar con otros", key="part_pregunta")

            col1, col2 = st.columns(2)
            with col1:
                volver = st.form_submit_button("Retroceder", use_container_width=True)
            with col2:
                avanzar = st.form_submit_button("Avanzar a experiencias", use_container_width=True)
            if volver:
                _goto_participant_stage(1)
            elif avanzar:
                if _validate_participant_stage2():
                    _goto_participant_stage(3)

    else:  # stage 3
        with st.form("form_participante_stage3", clear_on_submit=False):
            st.subheader("Así ordenas tus experiencias")
            if "part_exp_order" not in st.session_state:
                st.session_state.part_exp_order = experiencias.copy()
            try:
                from streamlit_sortables import sort_items  # type: ignore
                st.caption("Arrastra para ordenar según tu interés (arriba = más interés)")
                current_order = st.session_state.part_exp_order
                sorted_items = sort_items(current_order, direction="vertical", key="exp_sort")

                # Mantén el orden para la siguiente interacción y como resultado final
                st.session_state.part_exp_order = sorted_items
                order = sorted_items
            except Exception:
                st.caption("Selecciona en orden de interés (sin repetir).")

                def ranker(options):
                    remaining = options.copy()
                    selected = []
                    for i in range(len(options)):
                        choice = st.selectbox(f"Puesto {i+1}", remaining, key=f"rank_{i}")
                        selected.append(choice)
                        remaining = [o for o in remaining if o != choice]
                    return selected

                order = ranker(experiencias)
                st.session_state.part_exp_order = order

            ranks = {exp: order.index(exp) + 1 for exp in experiencias}
            experiencia_top = order[0] if order else ""

            st.text_area(
                "¿Por qué te interesa la experiencia que pusiste de primera?",
                max_chars=1000,
                help="Puedes usar hasta 1000 caracteres para contarnos tu motivación.",
                key="part_motivo"
            )

            st.markdown("#### Nivel de experticie")
            st.slider(
                "Mueve la barra para ubicarte",
                min_value=1,
                max_value=3,
                step=1,
                key="part_perfil_slider"
            )
            perfil_map = {1: "Curioso", 2: "Explorador", 3: "Protagonista"}
            seleccionado = st.session_state.get("part_perfil_slider", 1)
            st.markdown(
                """
                <div class=\"perfil-slider-labels\">
                    <span>⭐ Curioso</span>
                    <span>⭐⭐ Explorador</span>
                    <span>⭐⭐⭐ Protagonista</span>
                </div>
                """,
                unsafe_allow_html=True,
            )
            perfil_cerc = perfil_map[seleccionado]
            st.text_area(
                "¿Tienes alguna pregunta sobre esa experiencia?",
                key="part_preguntas_frec"
            )

            st.markdown("#### Acompañamiento")
            st.caption("Durante el encuentro de Claveriado 2026 tendremos distintas actividades de acompañamiento. Marca los acompañamientos con los que cuentas o quisieras fortalecer.")
            col_a, col_b, col_c = st.columns(3)
            col_d, col_e, col_f = st.columns(3)
            col_a.checkbox("Familia", key="part_acomp_familia")
            col_b.checkbox("Amigos", key="part_acomp_amigos")
            col_c.checkbox(
                "Escucha activa / apoyo emocional",
                key="part_acomp_escucha",
            )
            col_d.checkbox("Mentoría o tutoría", key="part_acomp_mentoria")
            col_e.checkbox("Acompañamiento espiritual", key="part_acomp_espiritual")
            col_f.checkbox(
                "Red comunitaria o institucional",
                key="part_acomp_red_comunidad",
            )
            st.checkbox("Ninguna por ahora", key="part_acomp_ninguna")

            current_none = bool(st.session_state.get("part_acomp_ninguna"))
            current_values = {
                key: bool(st.session_state.get(key))
                for key in ACOMP_KEYS
            }
            prev_values = st.session_state.get("_prev_part_acomp_values")
            prev_none = st.session_state.get("_prev_part_acomp_ninguna")

            none_toggled_on = current_none and not bool(prev_none)
            if prev_values is None:
                toggled_other_keys = [key for key, val in current_values.items() if val]
            else:
                toggled_other_keys = [
                    key
                    for key, val in current_values.items()
                    if val and not bool(prev_values.get(key))
                ]

            if none_toggled_on:
                for key in ACOMP_KEYS:
                    if st.session_state.get(key):
                        st.session_state[key] = False
            elif toggled_other_keys and current_none:
                st.session_state["part_acomp_ninguna"] = False

            st.session_state["_prev_part_acomp_values"] = current_values
            st.session_state["_prev_part_acomp_ninguna"] = current_none

            conoce_opciones = ["", "Sí", "No", "Más o menos"]
            current_conoce = st.session_state.get("part_conoce_rji", "")
            if current_conoce not in conoce_opciones:
                st.session_state.part_conoce_rji = ""
            st.selectbox(
                "¿Conoces qué es la RJI (Red Juvenil Ignaciana)?",
                conoce_opciones,
                key="part_conoce_rji",
                format_func=lambda val: "Selecciona una opción" if val == "" else val,
            )

            st.markdown("---")
            st.markdown(
                "La información recolectada es sensible y se utilizará únicamente para construir tu perfil en la Claveriada RJI y para la logística del encuentro. Consulta la [política de tratamiento de datos personales](https://jesuitas.co/wp-content/uploads/2023/08/politica-de-tratamiento-de-datos-personales.pdf)."
            )
            st.checkbox(
                "Acepto el tratamiento de datos personales conforme a la política indicada",
                key="part_acepta_datos"
            )
            st.checkbox(
                "Autorizo recibir información y contacto directo de la RJI vía WhatsApp",
                key="part_acepta_whatsapp"
            )

            col_back, col_save = st.columns(2)
            with col_back:
                volver_etapa = st.form_submit_button("Retroceder", use_container_width=True)
            with col_save:
                guardar = st.form_submit_button("Guardar participante", use_container_width=True)
            if volver_etapa:
                _goto_participant_stage(2)
            elif guardar:
                if not _validate_participant_stage3():
                    pass
                else:
                    es_mayor = st.session_state.get("part_es_mayor_option") == "Sí"
                    doc_p_clean = st.session_state.get("_clean_part_doc_p", "").strip()
                    if not doc_p_clean:
                        doc_ok, normalized = _normalize_numeric_input(st.session_state.get("part_doc_p", ""))
                        doc_p_clean = normalized if doc_ok else st.session_state.get("part_doc_p", "").strip()
                    doc_p = doc_p_clean

                    doc_a_clean = st.session_state.get("_clean_part_doc_a", "").strip()
                    if not doc_a_clean:
                        doc_a_ok, normalized_a = _normalize_numeric_input(st.session_state.get("part_doc_a", ""))
                        doc_a_clean = normalized_a if doc_a_ok else st.session_state.get("part_doc_a", "").strip()
                    doc_a = doc_a_clean

                    nom_a = st.session_state.get("part_nom_a", "")

                    ts = datetime.now(ZoneInfo("America/Bogota")).isoformat(timespec="seconds")
                    intereses = st.session_state.get("part_intereses", [])
                    conoce_map = {"Sí": "Si", "No": "No", "Más o menos": "Mas o menos", "": ""}
                    acomp_items = []
                    if st.session_state.get("part_acomp_familia"):
                        acomp_items.append("Familia")
                    if st.session_state.get("part_acomp_amigos"):
                        acomp_items.append("Amigos")
                    if st.session_state.get("part_acomp_escucha"):
                        acomp_items.append("Escucha activa / apoyo emocional")
                    if st.session_state.get("part_acomp_mentoria"):
                        acomp_items.append("Mentoría o tutoría")
                    if st.session_state.get("part_acomp_espiritual"):
                        acomp_items.append("Acompañamiento espiritual")
                    if st.session_state.get("part_acomp_red_comunidad"):
                        acomp_items.append("Red comunitaria o institucional")
                    if st.session_state.get("part_acomp_ninguna"):
                        acomp_items.append("Ninguna")

                    payload = _get_participant_payload()

                    def _capture_field(
                        key: str,
                        raw_value: object,
                        *,
                        sanitizer=_clean_string,
                        allow_empty: bool = False,
                    ) -> str:
                        existing = payload.get(key, "")
                        clean = sanitizer(raw_value)
                        if clean:
                            payload[key] = clean
                            return clean
                        if allow_empty:
                            if key not in payload:
                                payload[key] = ""
                                return ""
                            return payload.get(key, "")
                        return existing

                    if doc_p:
                        payload["documento_participante"] = doc_p
                    if st.session_state.get("part_tipo_doc_p") or not payload.get("tipo_documento_participante"):
                        payload["tipo_documento_participante"] = st.session_state.get("part_tipo_doc_p", "")
                    payload["es_mayor_edad"] = es_mayor

                    if doc_a or "documento_contacto" not in payload:
                        payload["documento_contacto"] = doc_a
                    if st.session_state.get("part_tipo_doc_a") or not payload.get("tipo_documento_contacto"):
                        payload["tipo_documento_contacto"] = st.session_state.get("part_tipo_doc_a", "")

                    uploads_dir = Path("uploads")
                    participante_doc_url = payload.get("archivo_doc_participante", "")
                    participante_label = payload.get("archivo_doc_participante_label", "")
                    participant_drive_failed = False

                    if st.session_state.get("part_doc_id_bytes") and st.session_state.get("part_doc_id_name"):
                        uploads_dir.mkdir(exist_ok=True)
                        participante_filename = f"{doc_p}_{st.session_state['part_doc_id_name']}"
                        participante_path = uploads_dir / participante_filename
                        with open(participante_path, "wb") as f:
                            f.write(st.session_state["part_doc_id_bytes"])
                        participante_label = participante_path.name
                        drive_link = ""
                        hasher = hashlib.sha256()
                        hasher.update(st.session_state["part_doc_id_bytes"])
                        hasher.update(f"|{UPLOADS_DRIVE_FOLDER_ID}".encode("utf-8"))
                        part_hash = hasher.hexdigest()
                        cached_hash = st.session_state.get("_part_doc_drive_hash")
                        cached_link = st.session_state.get("_part_doc_drive_link")
                        if cached_hash == part_hash and cached_link:
                            drive_link = cached_link
                        else:
                            drive_link = upload_file_to_drive(participante_path, UPLOADS_DRIVE_FOLDER_ID)
                            if drive_link:
                                st.session_state["_part_doc_drive_hash"] = part_hash
                                st.session_state["_part_doc_drive_link"] = drive_link
                        if drive_link:
                            participante_doc_url = drive_link
                        else:
                            participante_doc_url = str(participante_path)
                            if UPLOADS_DRIVE_FOLDER_ID:
                                participant_drive_failed = True


                    nombres_val = _capture_field("nombres", st.session_state.get("part_nombres", ""))
                    apellidos_val = _capture_field("apellidos", st.session_state.get("part_apellidos", ""))
                    apodo_val = _capture_field(
                        "como_te_gusta_que_te_digan",
                        st.session_state.get("part_apodo", ""),
                        allow_empty=True,
                    )
                    tel_val = _capture_field(
                        "telefono_celular",
                        st.session_state.get("part_tel", ""),
                        sanitizer=_clean_phone_number,
                    )
                    tel_val_sheet = _format_phone_for_sheet(tel_val)
                    correo_val = _capture_field("correo", st.session_state.get("part_correo", ""))
                    direccion_val = _capture_field("direccion", st.session_state.get("part_direccion", ""))
                    region_val = _capture_field("region", st.session_state.get("part_region", ""))
                    ciudad_val = _capture_field("ciudad", st.session_state.get("part_ciudad", ""))
                    eps_val = _capture_field(
                        "eps",
                        st.session_state.get("part_eps", ""),
                        allow_empty=True,
                    )
                    rest_alim_val = _capture_field(
                        "restricciones_alimentarias",
                        st.session_state.get("part_rest_alim", ""),
                        allow_empty=True,
                    )
                    salud_mental_val = _capture_field(
                        "salud_mental",
                        st.session_state.get("part_salud_mental", ""),
                        allow_empty=True,
                    )
                    obra_val = _capture_field("obra_institucion", st.session_state.get("part_obra", ""))
                    proceso_val = _capture_field(
                        "proceso_juvenil",
                        st.session_state.get("part_proceso", ""),
                        allow_empty=True,
                    )

                    exp_sig_val = payload.get("experiencia_significativa") or _clean_string(st.session_state.get("part_exp_sig", ""))
                    if exp_sig_val:
                        payload["experiencia_significativa"] = exp_sig_val
                    intereses_payload = payload.get("intereses_personales")
                    if intereses:
                        payload["intereses_personales"] = list(intereses)
                        intereses_payload = payload["intereses_personales"]
                    dato_freak_val = _capture_field(
                        "hobby_o_dato_curioso",
                        st.session_state.get("part_dato_freak", ""),
                        allow_empty=True,
                    )
                    pregunta_val = _capture_field(
                        "pregunta_para_conectar",
                        st.session_state.get("part_pregunta", ""),
                        allow_empty=True,
                    )
                    motivo_val = _capture_field("motivo_experiencia_top", st.session_state.get("part_motivo", ""))
                    preguntas_frec_val = _capture_field(
                        "preguntas_frecuentes",
                        st.session_state.get("part_preguntas_frec", ""),
                        allow_empty=True,
                    )

                    nom_a_clean = _capture_field("nombres_contacto", nom_a)
                    ape_a_clean = _capture_field("apellidos_contacto", st.session_state.get("part_ape_a", ""))
                    tel_a_clean = _capture_field(
                        "telefono_contacto",
                        st.session_state.get("part_tel_a", ""),
                        sanitizer=_clean_phone_number,
                    )
                    tel_a_sheet = _format_phone_for_sheet(tel_a_clean)
                    correo_a_clean = _capture_field(
                        "correo_contacto",
                        st.session_state.get("part_correo_a", ""),
                        allow_empty=True,
                    )
                    parentesco_clean = _capture_field("parentesco_contacto", st.session_state.get("part_parentesco_a", ""))

                    fecha_nac_value = st.session_state.get("part_fecha_nac")
                    if fecha_nac_value is not None or "fecha_nacimiento" not in payload:
                        payload["fecha_nacimiento"] = fecha_nac_value
                    talla_value = st.session_state.get("part_talla", "")
                    if talla_value or "talla_camisa" not in payload:
                        payload["talla_camisa"] = talla_value

                    payload["acompanamientos_marcados"] = ", ".join(acomp_items)
                    payload["acompanamiento_familia"] = bool(st.session_state.get("part_acomp_familia"))
                    payload["acompanamiento_amigos"] = bool(st.session_state.get("part_acomp_amigos"))
                    payload["acompanamiento_escucha_activa"] = bool(st.session_state.get("part_acomp_escucha"))
                    payload["acompanamiento_mentoria"] = bool(st.session_state.get("part_acomp_mentoria"))
                    payload["acompanamiento_espiritual"] = bool(st.session_state.get("part_acomp_espiritual"))
                    payload["acompanamiento_red_comunitaria"] = bool(st.session_state.get("part_acomp_red_comunidad"))
                    payload["acompanamiento_ninguna"] = bool(st.session_state.get("part_acomp_ninguna"))

                    conoce_value = conoce_map.get(st.session_state.get("part_conoce_rji"), "")
                    if conoce_value or "conoce_rji" not in payload:
                        payload["conoce_rji"] = conoce_value

                    acepta_datos = bool(st.session_state.get("part_acepta_datos"))
                    acepta_whatsapp = bool(st.session_state.get("part_acepta_whatsapp"))
                    payload["acepta_tratamiento_datos"] = acepta_datos
                    payload["acepta_whatsapp"] = acepta_whatsapp
                    payload["experiencia_top_calculada"] = experiencia_top
                    payload["nivel_experticie"] = perfil_cerc
                    if participante_doc_url:
                        payload["archivo_doc_participante"] = participante_doc_url
                    elif "archivo_doc_participante" not in payload:
                        payload["archivo_doc_participante"] = ""
                    if participante_label:
                        payload["archivo_doc_participante_label"] = participante_label
                    elif "archivo_doc_participante_label" not in payload:
                        payload["archivo_doc_participante_label"] = ""
                    full_name = f"{nombres_val} {apellidos_val}".strip()
                    edad_aprox = calcular_edad(payload.get("fecha_nacimiento"))
                    intereses_text = ", ".join(intereses_payload or [])
                    participante_doc_cell = _format_upload_for_sheet(
                        payload.get("archivo_doc_participante", participante_doc_url),
                        payload.get("archivo_doc_participante_label", participante_label),
                    )
                    drive_error_message = st.session_state.get("_drive_last_error", "").strip()
                    if drive_error_message and participant_drive_failed:
                        st.warning(
                            "No se pudo publicar uno o más archivos en Drive. Se guardó la ruta local por ahora. "
                            "Mensaje técnico: " + drive_error_message
                        )

                    row = [
                        ts,
                        "TRUE" if es_mayor else "FALSE",
                        payload.get("tipo_documento_participante", ""),
                        payload.get("documento_participante", doc_p),
                        nombres_val,
                        apellidos_val,
                        full_name,
                        apodo_val,
                        tel_val_sheet,
                        correo_val,
                        direccion_val,
                        region_val,
                        ciudad_val,
                        str(payload.get("fecha_nacimiento")),
                        edad_aprox,
                        payload.get("talla_camisa", ""),
                        eps_val,
                        rest_alim_val,
                        salud_mental_val,
                        obra_val,
                        proceso_val,
                        intereses_text,
                        exp_sig_val,
                        dato_freak_val,
                        pregunta_val,
                        int(ranks["Servicio"]),
                        int(ranks["Peregrinaje"]),
                        int(ranks["Cultura y arte"]),
                        int(ranks["Espiritualidad"]),
                        int(ranks["Vocación"]),
                        int(ranks["Incidencia política"]),
                        experiencia_top,
                        perfil_cerc,
                        motivo_val,
                        preguntas_frec_val,
                        payload.get("acompanamientos_marcados", ", ".join(acomp_items)),
                        "TRUE" if payload.get("acompanamiento_familia") else "FALSE",
                        "TRUE" if payload.get("acompanamiento_amigos") else "FALSE",
                        "TRUE" if payload.get("acompanamiento_escucha_activa") else "FALSE",
                        "TRUE" if payload.get("acompanamiento_mentoria") else "FALSE",
                        "TRUE" if payload.get("acompanamiento_espiritual") else "FALSE",
                        "TRUE" if payload.get("acompanamiento_red_comunitaria") else "FALSE",
                        "TRUE" if payload.get("acompanamiento_ninguna") else "FALSE",
                        payload.get("conoce_rji", conoce_value),
                        payload.get("tipo_documento_contacto", ""),
                        payload.get("documento_contacto", ""),
                        nom_a_clean,
                        ape_a_clean,
                        tel_a_sheet,
                        correo_a_clean,
                        parentesco_clean,
                        participante_doc_cell,
                        "TRUE" if acepta_datos else "FALSE",
                        "TRUE" if acepta_whatsapp else "FALSE",
                    ]
                    try:
                        append_row(SPREADSHEET_ID, SHEET_NAME, row, PARTICIPANTES_COLS)
                        try:
                            update_unificado(SPREADSHEET_ID)
                        except Exception:
                            pass
                        st.session_state["_participant_success_message"] = "¡Tu registro quedó guardado! Gracias por llegar al final ✨"
                        st.session_state["_participant_reset_pending"] = True
                        st.experimental_rerun()
                    except Exception as e:
                        st.error(f"No se pudo guardar: {e}")

    render_stage_progress(stage)

# ================= ACOMPAÑANTE =================
with tab2:
    st.info("RED JUVENIL IGNACIANA ESTÁ EN PROCESO DE SELECCIÓN DE ACOMPAÑANTES.")

# ================= VOLUNTARIADO =================
with tab3:
    st.info("RED JUVENIL IGNACIANA ESTÁ EN PROCESO DE SELECCIÓN DE VOLUNTARIOS.")

st.markdown("</div>", unsafe_allow_html=True)
