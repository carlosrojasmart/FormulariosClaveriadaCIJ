# doc.py
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date, datetime
from pathlib import Path
import pandas as pd
import gspread
from google.oauth2.service_account import Credentials

# ====== CONFIGURA AQUÍ ======
INFO_EVENTO = {
    "nombre": "Claveriada RJI",
    "fecha": "Por Confirmar",
    "lugar": "Por Confirmar",
    "descripcion": (
        "El Encuentro Juvenil RJI es un espacio formativo y de convivencia. "
        "Durante las actividades se promueven valores, el cuidado integral y el acompañamiento a los/las jóvenes. "
        "El acompañante se compromete a apoyar los protocolos de bienestar, seguridad y convivencia."
    ),
}
LOGO_PATH = "assets/logo.png"  


# ====== UTILIDADES ======
def calcular_edad(desde_fecha_str):
    """Calcula edad aproximada desde una fecha 'YYYY-MM-DD' o similar."""
    if not desde_fecha_str:
        return ""
    try:
        # aceptar varios formatos
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%Y/%m/%d", "%d-%m-%Y"):
            try:
                d = datetime.strptime(str(desde_fecha_str), fmt).date()
                break
            except ValueError:
                d = None
        if d is None:
            # último intento: pandas parse
            d = pd.to_datetime(desde_fecha_str, errors="coerce").date()
        today = date.today()
        return today.year - d.year - ((today.month, today.day) < (d.month, d.day))
    except Exception:
        return ""


def estilo_titulo(p, texto, size=16, bold=True, align="center"):
    p.clear()
    run = p.add_run(texto)
    run.bold = bold
    run.font.size = Pt(size)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_espacio(doc, alto_pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(alto_pts)


def set_margenes(doc, cm_top=2.0, cm_bottom=2.0, cm_left=2.0, cm_right=2.0):
    section = doc.sections[0]
    section.top_margin = Cm(cm_top)
    section.bottom_margin = Cm(cm_bottom)
    section.left_margin = Cm(cm_left)
    section.right_margin = Cm(cm_right)


def celda_texto(celda, texto, bold=False, size=11):
    celda.text = ""
    p = celda.paragraphs[0]
    run = p.add_run(str(texto))
    run.bold = bold
    run.font.size = Pt(size)


def crear_doc_consentimiento(
    nombre_acomp, doc_acomp, correo_acomp="", tel_acomp="",
    participantes=None,  # lista de dicts: [{nombre, documento, fecha_nacimiento, eps, salud}, ...]
    info_evento=INFO_EVENTO,
    logo_path=LOGO_PATH
):
    doc = Document()
    set_margenes(doc, 2, 2, 2, 2)

    # Encabezado con logo + título
    if logo_path and Path(logo_path).exists():
        header = doc.sections[0].header
        hdr_p = header.paragraphs[0]
        run = hdr_p.add_run()
        try:
            run.add_picture(logo_path, width=Inches(3.0))
        except Exception:
            pass

    titulo = doc.add_paragraph()
    estilo_titulo(titulo, "FORMATO DE AUTORIZACIÓN Y ACOMPAÑAMIENTO", size=16, bold=True, align="center")
    sub = doc.add_paragraph()
    estilo_titulo(sub, info_evento.get("nombre", "Encuentro RJI"), size=13, bold=False, align="center")

    add_espacio(doc, 6)

    # Información del Encuentro
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_info.add_run(
        f"Fecha: {info_evento.get('fecha','')} • Lugar: {info_evento.get('lugar','')}\n"
        f"{info_evento.get('descripcion','')}"
    )
    run.font.size = Pt(10)

    add_espacio(doc, 6)

    # Datos del acompañante
    p_datos = doc.add_paragraph()
    p_datos.style = doc.styles["Normal"]
    run = p_datos.add_run(
        f"Yo, {nombre_acomp}, identificado(a) con documento No. {doc_acomp}, "
        f"en calidad de acudiente/acompañante, autorizo la participación de las/los siguientes jóvenes "
        f"en el {info_evento.get('nombre','Encuentro')}."
    )
    run.font.size = Pt(11)

    add_espacio(doc, 4)

    # Contacto acompañante (bloque pequeño)
    t_acomp = doc.add_table(rows=2, cols=2)
    t_acomp.style = "Table Grid"
    t_acomp.autofit = True
    celda_texto(t_acomp.cell(0, 0), "Correo del acompañante", bold=True)
    celda_texto(t_acomp.cell(0, 1), correo_acomp or "_______________________________")
    celda_texto(t_acomp.cell(1, 0), "Teléfono del acompañante", bold=True)
    celda_texto(t_acomp.cell(1, 1), tel_acomp or "_______________________________")

    add_espacio(doc, 8)

    # Tabla de jóvenes: sin firma del joven, con Edad, EPS, Complicaciones de salud
    doc.add_paragraph().add_run("Relación de jóvenes a cargo").bold = True
    tabla = doc.add_table(rows=1, cols=5)
    tabla.style = "Table Grid"
    hdr = tabla.rows[0].cells
    celda_texto(hdr[0], "Nombre completo", bold=True)
    celda_texto(hdr[1], "Documento", bold=True)
    celda_texto(hdr[2], "Edad", bold=True)
    celda_texto(hdr[3], "EPS", bold=True)
    celda_texto(hdr[4], "Complicaciones de salud", bold=True)

    if participantes:
        for it in participantes:
            nombre = it.get("nombre", "")
            documento = it.get("documento", "")
            fnac = it.get("fecha_nacimiento", "")
            edad = calcular_edad(fnac)
            eps = it.get("eps", "")
            # “Complicaciones de salud”: puedes combinar alergias / restricciones / salud_mental
            salud = it.get("salud", "")
            fila = tabla.add_row().cells
            celda_texto(fila[0], nombre)
            celda_texto(fila[1], documento)
            celda_texto(fila[2], edad)
            celda_texto(fila[3], eps)
            celda_texto(fila[4], salud)

    add_espacio(doc, 10)

    # Compromiso / texto legal breve
    p_comp = doc.add_paragraph(
        "Declaro que la información consignada es veraz. Me comprometo a acompañar y velar por el bienestar de las/los jóvenes, "
        "cumplir las indicaciones del equipo organizador y notificar cualquier situación de salud o emergencia."
    )
    p_comp_format = p_comp.paragraph_format
    p_comp_format.space_after = Pt(6)

    # Firmas lado a lado: acompañante (izq) / institución (der)
    firmas = doc.add_table(rows=2, cols=2)
    firmas.autofit = True
    firmas.cell(0, 0).width = Inches(3.5)
    firmas.cell(0, 1).width = Inches(3.5)

    # Líneas de firma
    celda_texto(firmas.cell(0, 0), "\n\n_______________________________")
    celda_texto(firmas.cell(0, 1), "\n\n_______________________________")

    # Etiquetas
    celda_texto(firmas.cell(1, 0), "Firma del acompañante", bold=True)
    celda_texto(firmas.cell(1, 1), "Firma de la institución", bold=True)

    return doc


# ====== MODO 1: GENERAR DESDE LISTA MANUAL ======
def demo_manual():
    """Ejemplo manual: edita esta lista para probar rápidamente."""
    participantes = [
        {
            "nombre": "Juan Pérez",
            "documento": "10203040",
            "fecha_nacimiento": "2009-07-12",
            "eps": "SURA",
            "salud": "Asma leve. Alergia a mariscos.",
        },
        {
            "nombre": "María Gómez",
            "documento": "11223344",
            "fecha_nacimiento": "2008-11-03",
            "eps": "Compensar",
            "salud": "Ninguna reportada.",
        },
    ]
    doc = crear_doc_consentimiento(
        nombre_acomp="Carlos Rodríguez",
        doc_acomp="99887766",
        correo_acomp="carlos@ejemplo.com",
        tel_acomp="+57 300 123 4567",
        participantes=participantes,
    )
    out = "consentimiento_rji_demo.docx"
    doc.save(out)
    print(f"Documento generado: {out}")


# ====== MODO 2: GENERAR DESDE GOOGLE SHEETS ======
def generar_desde_google_sheet(
    spreadsheet_id,
    documento_acompanante,
    out_path=None,
    credentials_json_path=None,
    credentials_info=None,
):
    """
    Lee la pestaña PARTICIPANTES desde una hoja de Google y arma la lista para el acompañante dado.
    Debes pasar el ID de la hoja y las credenciales del Service Account (ruta al JSON o el dict ya cargado).
    """
    scopes = ["https://www.googleapis.com/auth/spreadsheets.readonly"]
    if credentials_info:
        creds = Credentials.from_service_account_info(credentials_info, scopes=scopes)
    elif credentials_json_path:
        creds = Credentials.from_service_account_file(credentials_json_path, scopes=scopes)
    else:
        raise ValueError("Debes proporcionar credentials_info o credentials_json_path para acceder a la hoja de cálculo.")

    client = gspread.authorize(creds)
    sh = client.open_by_key(spreadsheet_id)
    ws = sh.worksheet("PARTICIPANTES")
    registros = ws.get_all_records()
    dfp = pd.DataFrame(registros)
    if dfp.empty:
        dfp = pd.DataFrame(columns=[
            "es_mayor_edad","documento_contacto","nombre_completo","documento_participante",
            "fecha_nacimiento","eps","salud_mental","restricciones_alimentarias"
        ])
    for col in [
        "es_mayor_edad","documento_contacto","nombre_completo","documento_participante",
        "fecha_nacimiento","eps","salud_mental","restricciones_alimentarias"
    ]:
        if col not in dfp.columns:
            dfp[col] = ""
    # Filtrar menores que declararon a este acudiente
    m = (
        (dfp["es_mayor_edad"].astype(str).str.lower().isin(["false", "no", "0"]))
        & (
            dfp["documento_contacto"].astype(str).str.replace(r"\s+", "", regex=True)
            == str(documento_acompanante).replace(" ", "")
        )
    )
    sub = dfp[m].copy()

    participantes = []
    for _, r in sub.iterrows():
        salud_merge = ""
        if str(r.get("salud_mental", "")).strip():
            salud_merge += f"Salud: {r['salud_mental']}. "
        if str(r.get("restricciones_alimentarias", "")).strip() and r["restricciones_alimentarias"].strip().lower() not in ["ninguna", "no", ""]:
            salud_merge += f"Alimentación: {r['restricciones_alimentarias']}."
        participantes.append(
            {
                "nombre": r.get("nombre_completo", ""),
                "documento": r.get("documento_participante", ""),
                "fecha_nacimiento": r.get("fecha_nacimiento", ""),
                "eps": r.get("eps", ""),
                "salud": salud_merge.strip(),
            }
        )

    # Datos básicos del acompañante (si los tienes en otro sitio, pásalos por parámetro)
    nombre_acomp = "Acompañante"
    correo = ""
    tel = ""

    doc = crear_doc_consentimiento(
        nombre_acomp=nombre_acomp,
        doc_acomp=str(documento_acompanante),
        correo_acomp=correo,
        tel_acomp=tel,
        participantes=participantes,
    )

    if out_path is None:
        out_path = f"consentimiento_rji_{documento_acompanante}.docx"
    doc.save(out_path)
    print(f"Documento generado: {out_path}")


if __name__ == "__main__":
    # === Elige cómo probar ===
    # 1) Demo manual (rápido):
    demo_manual()

    # 2) Desde Google Sheets:
    # generar_desde_google_sheet(
    #     "1vVyWpY1izl8wcESXK-mdXNOqHzyc-yGUZcZ1H-0sMNs",
    #     documento_acompanante="99887766",
    #     credentials_json_path="ruta/a/tu-service-account.json",
    # )

   
